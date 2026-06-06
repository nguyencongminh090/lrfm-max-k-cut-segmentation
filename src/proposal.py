"""Generate the Research Proposal (.docx) following Class_Group_Research_Proposal_Template.

Populated with THIS project: an LRFM extension of the graph-based max-k-cut RFM
segmentation method, on Online Retail II. Content reflects what we actually did,
including the honest caveats (silhouette trade-off, no response data, +1.5% monetary).

Output: results/Research_Proposal_DAP391m.docx
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import config

OUT = config.RESULTS / "Research_Proposal_DAP391m.docx"
ACCENT = RGBColor(0x1F, 0x4E, 0x78)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def body(doc, text, italic=False):
    p = doc.add_paragraph(text)
    if italic:
        for r in p.runs:
            r.italic = True
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def build():
    doc = Document()
    # ----- title -----
    t = doc.add_heading("RESEARCH PROPOSAL", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("An LRFM Extension of Graph-Based Max-k-cut Customer "
                            "Segmentation for Targeted E-commerce Marketing")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.bold = True; r.font.size = Pt(13)

    # ----- header info table -----
    tbl = doc.add_table(rows=6, cols=2); tbl.style = "Table Grid"
    rows = [
        ("Research title", "An LRFM Extension of Graph-Based Max-k-cut Customer "
                           "Segmentation for Targeted E-commerce Marketing"),
        ("Course", "DAP391m"),
        ("Mentor", "Nguyễn Hoàng Linh"),
        ("Member 1", "Nguyễn Công Minh — SE203724"),
        ("Member 2", "Võ Đức Nhật — SE203702"),
        ("Member 3", "Trần Nguyễn Minh Hải — SE203718"),
    ]
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    # ----- Abstract -----
    h(doc, "Abstract", 1)
    body(doc,
        "Customer segmentation underpins customer relationship management, letting "
        "businesses target marketing to groups of similar customers. The Recency-"
        "Frequency-Monetary (RFM) model is the dominant behavioural segmentation "
        "framework, but it uses only three variables and is usually clustered with "
        "K-means, which offers no optimality guarantee. A recent graph-theoretic "
        "method reformulates RFM segmentation as the maximum k-cut problem and solves "
        "it via a vertex-reduction technique that is provably equivalent to clustering "
        "the full dataset, reportedly outperforming K-means and a Formal Concept "
        "Analysis baseline. This research asks whether augmenting that model with a "
        "Length (customer-tenure) dimension - extending RFM to LRFM - yields more "
        "actionable segments. We reproduce the graph-based baseline on the public "
        "Online Retail II dataset (5,878 customers) and extend it to four variables, "
        "preserving the vertex-reduction property (at most T^4 = 625 super-vertices). "
        "As no commercial optimisation license is available, we implement a from-"
        "scratch, license-free max-k-cut solver (multi-start local search with "
        "Kernighan-Lin refinement, JIT-compiled) validated against brute-force optima. "
        "Segments are ranked by a Cheng & Chen loyalty measure. The methodology "
        "combines data cleaning, a normalised (3NF) SQLite analytical layer with "
        "quantile scoring, graph construction and clustering, and a multi-method "
        "comparison. Expected results: faithful reproduction of the baseline, and "
        "LRFM segments that separate newly-acquired from long-tenured loyal customers "
        "- a distinction RFM cannot make - while honestly reporting that the added "
        "dimension trades the silhouette metric for managerial interpretability.")
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("Customer segmentation; RFM/LRFM model; Graph-based clustering; "
               "Maximum k-cut; Customer loyalty; E-commerce.")

    # ----- 1. Introduction -----
    h(doc, "1. Introduction", 1)

    h(doc, "1.1. Literature review", 2)
    body(doc,
        "Internationally, RFM analysis is the most widely used behavioural "
        "segmentation method, typically combined with clustering. Chen et al. (2012) "
        "applied RFM with K-means and decision trees to the online-retail dataset that "
        "has since become a standard benchmark, and Christy et al. (2021) proposed RFM "
        "ranking variants. K-means remains popular for its simplicity but is sensitive "
        "to initialisation and outliers and provides only a high-quality suboptimal "
        "solution. To improve interpretability, Rungruang et al. (2024) combined RFM "
        "with Formal Concept Analysis (FCA) to expose relationships among segments. "
        "Most recently, Correa Vianna Filho et al. (2026) modelled the problem as the "
        "maximum k-cut problem on a customer graph and introduced a vertex-reduction "
        "technique that makes an exact formulation tractable, outperforming K-means and "
        "the FCA approach by the silhouette index. Several extensions enrich RFM with a "
        "fourth variable - LRFM adds Length/tenure (Chang & Tsay, 2004), RFMT adds "
        "inter-purchase time (Zhou et al., 2021), and RFMTC adds a churn term (Yeh et "
        "al., 2009).")
    body(doc,
        "In Vietnam, RFM-based segmentation has been applied mainly in retail, "
        "e-commerce and banking CRM, usually with K-means, reflecting the same "
        "three-variable, distance-clustering paradigm seen internationally. "
        "(Specific Vietnamese references to be added by the authors as required.)")
    body(doc,
        "In practice these works support targeted marketing, customer-retention "
        "campaigns, loyalty programmes and customer-value ranking across retail, "
        "e-commerce and financial services.")

    h(doc, "1.2. The limitation of current works", 2)
    body(doc,
        "Existing approaches share several limitations. First, the RFM model uses only "
        "three variables and therefore cannot distinguish customers who share the same "
        "recency, frequency and monetary profile but differ in how long they have been "
        "active - e.g. a newly-acquired high-spender versus a long-tenured loyal "
        "customer. Second, K-means gives no optimality guarantee and is outlier-"
        "sensitive, while the FCA approach, though interpretable, scores lower on "
        "cluster quality. Third, the exact graph-based method depends on a commercial "
        "optimisation solver (Gurobi), whose free license is size-limited - a barrier "
        "to reproduction. Finally, segments are often labelled with arbitrary "
        "practitioner names rather than a principled, citable criterion.")

    h(doc, "1.3. The necessity of the research", 2)
    body(doc,
        "This research addresses the absence of a tenure/loyalty-duration dimension in "
        "graph-based RFM segmentation, and the practical barrier of solver licensing. "
        "Its originality lies in extending the recent max-k-cut formulation to LRFM "
        "while preserving the vertex-reduction guarantee, and in delivering a fully "
        "license-free, reproducible pipeline. The topic is timely given the growth of "
        "e-commerce and data-driven CRM. Scientifically, it contributes a faithful "
        "reproduction of a 2026 baseline, an honest evaluation of when an added "
        "dimension helps (managerial actionability) versus when it does not (the "
        "silhouette metric), and a critical observation that the baseline's reported "
        "advantage over K-means is partly an artefact of differing preprocessing.")

    # ----- Research objectives -----
    h(doc, "Research objectives", 1)
    for o in [
        "Reproduce the graph-based max-k-cut RFM segmentation baseline on the Online "
        "Retail II dataset and verify it against the published results.",
        "Extend the RFM model to LRFM by adding a Length (tenure) dimension and assess "
        "whether L provides non-redundant, actionable segmentation signal.",
        "Implement and validate a license-free max-k-cut solver (multi-start local "
        "search with Kernighan-Lin refinement) as an alternative to a commercial solver.",
        "Build a normalised (3NF) SQL analytical layer as the single source of truth "
        "for R/F/M/L scores and segment analytics.",
        "Compare the proposed method against K-means, hierarchical (Ward) and Gaussian "
        "Mixture clustering on identical features, and characterise segments by a "
        "Cheng & Chen loyalty ranking.",
    ]:
        numbered(doc, o)

    # ----- Research scope -----
    h(doc, "Research scope", 1)
    body(doc,
        "The study focuses on offline customer segmentation using the public Online "
        "Retail II dataset (UCI), comprising 5,878 customers after cleaning. It covers "
        "data preparation, RFM/LRFM scoring, graph-based clustering and segment "
        "profiling. It does not build a live CRM system, and it does not perform "
        "supervised response prediction because the dataset contains no campaign-"
        "response label; predictive modelling is therefore out of scope.")

    # ----- Feasibility -----
    h(doc, "Feasibility of research", 1)
    body(doc,
        "The project is highly feasible. The dataset is public and already obtained. "
        "The baseline has already been reproduced (5,878 customers; recency mean "
        "200.87; 114-vertex reduced graph; silhouette within 0.02 of the published "
        "values). The from-scratch solver has been validated against brute-force "
        "optima (40/40) and runs in seconds on the reduced graph (<= 625 vertices). "
        "All tools are open-source, so no licensing or specialised hardware is needed.")

    # ----- Approach and Method -----
    h(doc, "Approach and Method", 1)
    body(doc,
        "The research adopts a reproduction-and-extension, experimental approach. The "
        "pipeline architecture is:")
    for s in [
        "Data preparation: clean Online Retail II (remove missing IDs, non-positive "
        "quantity/price, duplicates; retain outliers per the baseline).",
        "Relational layer: load a 3NF SQLite schema (dim_customer, dim_product, "
        "dim_invoice, fact_line) and compute R/F/M/L quantile scores (NTILE) as the "
        "source of truth.",
        "Graph construction: represent each customer as a vertex with edges weighted by "
        "the Manhattan distance between (L)RFM score vectors; merge identical scores "
        "into super-vertices (<= T^q), preserving optimality.",
        "Clustering: solve the maximum k-cut on the reduced graph with a license-free "
        "multi-start local-search + Kernighan-Lin solver; lift the assignment back to "
        "all customers.",
        "Evaluation: sweep k = 2..10, compute the silhouette index, and compare against "
        "K-means, Ward and Gaussian Mixture on identical features.",
        "Segment analysis: rank clusters by Cheng & Chen loyalty distance and profile "
        "them; produce SQL analytics and modern visualisations.",
    ]:
        numbered(doc, s)
    body(doc,
        "Correctness is established by reproducing the baseline's reconciliation "
        "anchors and by validating the solver against brute-force optima.")

    # ----- Research plan -----
    h(doc, "Research plan", 1)
    plan = [
        ("1", "Week 1", "Literature review; baseline reproduction setup; data download",
         "Scope, reproduced RFM baseline", "Research team"),
        ("2", "Week 2", "Data cleaning, EDA, and 3NF SQLite database",
         "Cleaned dataset, retail.db, EDA report", "Research team"),
        ("3", "Week 3", "Feature engineering: LRFM scoring and justification",
         "customer_rfml table, LRFM justification", "Research team"),
        ("4", "Week 4", "License-free solver implementation & validation; modelling (RFM vs LRFM) & multi-method comparison",
         "Validated solver, test suite, and comparison/silhouette results", "Research team"),
        ("5", "Week 5", "SQL analytics, visualisation, report, proposal & audit log",
         "Figures, dashboards, final report", "Research team"),
    ]
    pt = doc.add_table(rows=1, cols=5); pt.style = "Table Grid"
    hdr = ["No.", "Date", "Task", "Output", "Person in charge"]
    for c, txt in zip(pt.rows[0].cells, hdr):
        c.text = txt; c.paragraphs[0].runs[0].bold = True
    for row in plan:
        cells = pt.add_row().cells
        for c, txt in zip(cells, row):
            c.text = txt

    # ----- Computational resources -----
    h(doc, "Computational Resource Requirements", 1)
    body(doc,
        "A standard laptop with no GPU. Python 3.13 with pandas, NumPy, scikit-learn, "
        "Numba, SQLite, and matplotlib/seaborn. No commercial optimisation solver is "
        "required: the max-k-cut solver is implemented from scratch and the reduced "
        "graph (<= 625 vertices) is solved in seconds.")

    # ----- Expected results -----
    h(doc, "Expected results", 1)
    body(doc,
        "The project is expected to deliver: (i) a faithful, reproducible "
        "implementation of the graph-based max-k-cut RFM baseline; (ii) an LRFM "
        "extension that isolates newly-acquired from long-tenured loyal customers - an "
        "actionable distinction RFM alone cannot make; (iii) an honest finding that "
        "adding the tenure dimension trades cluster-separation (silhouette) for "
        "managerial interpretability; and (iv) a critical result that, on identical "
        "features, max-k-cut and K-means perform comparably, indicating the baseline's "
        "reported advantage is partly a preprocessing artefact. Practically, the "
        "loyalty-ranked segments support differentiated retention, win-back and "
        "nurture campaigns, and the open, license-free pipeline is reusable for other "
        "retail datasets.")

    # ----- References -----
    h(doc, "References", 1)
    refs = [
        "Chang, H. H., & Tsay, S. F. (2004). Integrating SOM and K-means in data mining "
        "clustering: an empirical study of CRM and profitability evaluation. Journal of "
        "Information Management, 11(4), 161-203.",
        "Chen, D., Sain, S. L., & Guo, K. (2012). Data mining for the online retail "
        "industry: a case study of RFM model-based customer segmentation using data "
        "mining. Journal of Database Marketing & Customer Strategy Management, 19(3), 197-208.",
        "Cheng, C.-H., & Chen, Y.-S. (2009). Classifying the segmentation of customer "
        "value via RFM model and RS theory. Expert Systems with Applications, 36(3), 4176-4184.",
        "Christy, A. J., Umamakeswari, A., Priyatharsini, L., & Neyaa, A. (2021). RFM "
        "ranking - an effective approach to customer segmentation. Journal of King Saud "
        "University - Computer and Information Sciences, 33(10), 1251-1257.",
        "Correa Vianna Filho, A. L., de Lima, L., & Kleina, M. (2026). RFM model "
        "customer segmentation from a graph theory perspective. Quality & Quantity. "
        "https://doi.org/10.1007/s11135-026-02784-0 (preprint arXiv:2505.08136).",
        "Ma, F., & Hao, J.-K. (2017). A multiple search operator heuristic for the "
        "max-k-cut problem. Annals of Operations Research, 248, 365-403.",
        "Rousseeuw, P. J. (1987). Silhouettes: a graphical aid to the interpretation and "
        "validation of cluster analysis. Journal of Computational and Applied "
        "Mathematics, 20, 53-65.",
        "Rungruang, C., Riyapan, P., Intarasit, A., Chuarkham, K., & Muangprathub, J. "
        "(2024). RFM model customer segmentation based on hierarchical approach using "
        "FCA. Expert Systems with Applications, 237, 121449.",
        "van Dam, E. R., & Sotirov, R. (2016). New bounds for the max-k-cut and "
        "chromatic number of a graph. Linear Algebra and its Applications, 488, 216-234.",
        "Yeh, I.-C., Yang, K.-J., & Ting, T.-M. (2009). Knowledge discovery on RFM model "
        "using Bernoulli sequence. Expert Systems with Applications, 36(3), 5866-5871.",
        "Zhou, J., Wei, J., & Xu, B. (2021). Customer segmentation by web content "
        "mining. Journal of Retailing and Consumer Services, 61, 102588.",
        "UCI Machine Learning Repository (2019). Online Retail II Data Set. "
        "https://archive.ics.uci.edu/dataset/502/online+retail+ii",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    config.RESULTS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Proposal written: {OUT}")
    print(f"  sections + {len(refs)} references; abstract ~230 words")


if __name__ == "__main__":
    build()
